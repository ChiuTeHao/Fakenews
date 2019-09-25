#!/usr/bin/python
#-*- encoding: UTF-8 -*-
import pandas as pd
import pickle
import re
import numpy as np
import csv
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
import json
import os
import time as Time

DEBUG = False
FILTER_FEW_ELEMENT = True

# 異體字
variant_w = [('台','臺'),('羗','羌')]

    # filter 多餘 extract entries
def smaller(comp1, comp2, comb_idx):
    chk_idx = [i for i in range(len(comp1)) if i != comb_idx]
    diff = len([i for i in chk_idx if comp1[i]!=comp2[i]])
    # if (comp1[comb_idx] in comp2[comb_idx]) and (comp2[comb_idx] != 'None') and (comp1[comb_idx] != comp2[comb_idx]) and (diff == 1):
    print(comp1)
    print(comp2)
    print(comb_idx)
    if (comp1[comb_idx] in comp2[comb_idx]) and diff==0:    
        return True
    else:
        return False

def variant_word_handle(list_, pair_list):
    list_all = []
    for (w_1, w_2) in pair_list:
        list_1 = [x.replace(w_1, w_2) for x in list_]
        list_2 = [x.replace(w_2, w_1) for x in list_]
        list_all = list(set(list_all + list_1+ list_2))
    return list_all

def variant_word_handle_dict(dict_, pair_list):
    list_ = dict_.keys()
    change_0 = [x[0] for x in pair_list]
    change_1 = [x[1] for x in pair_list]
    ch_target_0 = [[x for x in list_ if i in x] for i in change_0]
    ch_target_1 = [[x for x in list_ if i in x] for i in change_1]
    for idx in range(len(pair_list)):
        for _it in ch_target_0[idx]:
            _it_new = _it.replace(pair_list[idx][0],pair_list[idx][1])
            if _it_new not in dict_.keys():
                dict_[_it_new] = dict_[_it]
    for idx in range(len(pair_list)):
        for _it in ch_target_1[idx]:
            _it_new = _it.replace(pair_list[idx][1],pair_list[idx][0])
            if _it_new not in dict_.keys():
                dict_[_it_new] = dict_[_it]
    return dict_

df = pd.read_csv('loc.csv')
county = list(set(df['County']))
county = [x for x in county if str(x)!='nan']
county = variant_word_handle(county, variant_w)
town = list(set(df['Town']))
town = [x for x in town if str(x)!='nan']
town = variant_word_handle(town, variant_w)
village = list(set(df['Village']))
village = [x for x in village if str(x)!='nan']
village = variant_word_handle(village, variant_w)
county_simp = []
for _it in county:
    county_simp.append(_it[:-1])
town_simp = []
for _it in town:
    if len(_it[:-1])>1 and _it[:-1] not in county_simp:
        town_simp.append(_it[:-1])
# print(len(county))
# print(len(town))
# print(len(village))

place_dict = dict()
for _it in county:
    place_dict[_it] = 'County'
for _it in county_simp:
    place_dict[_it] = 'County'
for _it in town:
    place_dict[_it] = 'Town'
for _it in town_simp:
    place_dict[_it] = 'Town'
for _it in village:
    place_dict[_it] = 'Village'
place_dict['五里埔'] = 'Village'
# print('五里埔' in place_dict.keys())
# print('那瑪夏' in place_dict.keys())

   

#road parse
# with open('road_name.csv','r',encoding='UTF-8') as fin:
#     lines = fin.readlines()
#     road_list = []
#     for line in lines:
#         a = line.strip().split(',')
#         if a[2][-1] not in {'巷','路','街','段'}:
#             a[2] += '路段'
#         if a[2] not in road_list:
#             road_list.append(a[2])
#     road_list = variant_word_handle(road_list, variant_w)
#     road_dict = dict()
#     for _it in road_list:
#         road_dict[_it] = 'Road'
# with open('data.p', 'wb') as f:
#     pickle.dump((place_dict, road_dict), f)

with open('road_name.csv','r',encoding='UTF-8') as fin:
    lines = fin.readlines()
    road_dict = dict()
    for line in lines:
        a = line.strip().split(',')
        if a[2][-1] not in {'巷','路','街','段'}:
            a[2] += '路路'
            if a[2] not in road_dict.keys():
                road_dict[a[2]] = 'Road_1'
        elif a[2] not in road_dict.keys():
            road_dict[a[2]] = 'Road'
    road_dict = variant_word_handle_dict(road_dict, variant_w)
    # road_dict = dict()
    # for _it in road_list:
    #     road_dict[_it] = 'Road'
with open('data.p', 'wb') as f:
    pickle.dump((place_dict, road_dict), f)
# print(road_dict['羌光路'])
#專有名詞
def filter(str_list, word_dict_list):
    str_list_1 = str_list[:]
    for _dict in word_dict_list:
        dict_list = list(_dict.keys())
        dict_list = sorted(dict_list, key = lambda x:len(x), reverse=True)        
        replace_list = []
        for term in dict_list:            
            idx = 0
            # print(str_list_1)
            while idx < len(str_list_1):
                string = str_list_1[idx]                
                if string[1] == 'S' and (term in string[0]):                    
                    # print(term,'a',  string[0])
                    str1 = string[0]
                    str1 = str1.replace(term,'|'+str(len(replace_list))+'|')
                    replace_list.append(term)                    
                    # print(str1)
                    # print(replace_list)
                    str_cut = re.split(r'(\|[0-9]+\|)', str1)
                    # print(str_cut)
                    for idx_s in range(len(str_cut)):
                        if re.match(r'\|[0-9]+\|',str_cut[idx_s]) is not None:
                            # print('aa', re.match('|[0-9]+|',str_cut[idx_s]))
                            _t = replace_list[int(str_cut[idx_s][1:-1])]
                            str_cut[idx_s] = (_t, _dict[_t])
                        else:
                            str_cut[idx_s] = (str_cut[idx_s], 'S')
                    del str_list_1[idx]
                    str_list_1 = str_list_1[:idx] + str_cut + str_list_1[idx:]
                    # for i in range(len(str_cut)):
                    #     str_list_1.insert(idx+i, str_cut[i])
                    idx += len(str_cut)
                else:
                    idx += 1
    # idx = 0 
    # while idx < str_list_1(str_list_1):
        
    idx = 0
    while idx < len(str_list_1):
        if len(str_list_1[idx][0])==0:
            # print(str_list_1[idx])
            del str_list_1[idx]
        else:
            idx += 1
    return str_list_1

# filter('',[road_dict])
#regular expression
def filter_re(str_list, re_rule, POS, target_POS = 'S', replace_re_rule = None):
    str_list_1 = str_list[:]
    # print('g11',str_list_1)
    idx = 0
    while idx < len(str_list_1):
        string = str_list_1[idx][0]
        # print('g11',str_list_1[idx])
        if str_list_1[idx][1] == target_POS and re.search(re_rule,string) is not None:
            cut_str = re.split('('+re_rule+')', string)
            for _idx, seg in enumerate(cut_str):
                if re.match(re_rule, seg) is not None:
                    cut_str[_idx] = (seg, 'Spec')
                else:
                    cut_str[_idx] = (seg, 'S')
            del str_list_1[idx]
            str_list_1 = str_list_1[:idx] + cut_str + str_list_1[idx:]
            idx += len(cut_str)
        else:
            idx += 1
    idx = 0
    while idx < len(str_list_1):        
        if len(str_list_1[idx][0])==0:
            # print(str_list_1[idx])
            del str_list_1[idx]
        else:
            idx += 1
    pre_POS = str_list_1[0][1]
    idx = 1
    while idx < len(str_list_1):
        if str_list_1[idx][1]==pre_POS and pre_POS != 'Spec':
            # print(str_list_1[idx])
            str_list_1[idx-1] = (str_list_1[idx-1][0] + str_list_1[idx][0], pre_POS)
            del str_list_1[idx]
        else:
            pre_POS = str_list_1[idx][1]
            idx += 1
    if replace_re_rule is not None:
        idx = 0
        while idx < len(str_list_1):
            if str_list_1[idx][1] == 'Spec':            
                for (pattern, repl) in replace_re_rule:
                    str_list_1[idx] = (re.sub(pattern, repl, str_list_1[idx][0]), 'Spec')
            idx += 1
    idx = 0
    while idx < len(str_list_1):
        if str_list_1[idx][1] == 'Spec':
            str_list_1[idx] = (str_list_1[idx][0], POS)
            # print('gg',str_list_1[idx])
        idx += 1
    return str_list_1

#for address parse
def filter_pre_re(str_list, pre_POS_list, re_rule, POS):
    str_list_1 = str_list[:]
    idx = 1
    while idx < len(str_list_1):
        # print('debuggg', str_list_1[idx-1][0],str_list_1[idx-1][1], str_list_1[idx-1][1] in pre_POS_list, str_list_1[idx][1]=='S', re.search(re_rule, str_list_1[idx][0]) is not None)
        if str_list_1[idx-1][1] in pre_POS_list and str_list_1[idx][1]=='S' and re.search(re_rule, str_list_1[idx][0]) is not None:
            str_sub_list = [str_list_1[idx]]
            # print('qqq', str_sub_list)
            str_sub_list = filter_re(str_sub_list, re_rule, POS)
            del str_list_1[idx]
            str_list_1 = str_list_1[:idx] + str_sub_list + str_list_1[idx:]
            idx += len(str_sub_list)
        else:
            idx += 1
    return str_list_1

def filter_re_2(str_list, re_rule, POS_list):
    str_list_1 = str_list[:]
    idx = 0
    while idx < len(str_list_1):
        idx_str, idx_POS = str_list_1[idx]         
        if idx_POS == 'S' and re.search(re_rule, idx_str) is not None:
            del str_list_1[idx]            
            cut_str = []
            start_idx = 0
            for match in re.finditer(re_rule, idx_str):
                match_idx = match.span()
                cut_str.append((str_list_1[start_idx:match_idx[0]],'S'))
                
# merge some of related POS, such as adress can be [county + town + village], time can be [month + date], ...
def merge_term(str_list, related_list, split_by_space = None):
    str_list_1 = str_list[:]
    buff = ['' for i in range(len(related_list))]
    for idx, (string, POS) in enumerate(str_list_1):
        if POS in related_list:
            buff[related_list.index(POS)] = string
            for i in range(related_list.index(POS)+1, len(related_list)):
                buff[i] = ''
            str_ = ''
            for i in range(related_list.index(POS),-1,-1):
                if buff[i] not in str_:
                    if split_by_space is not None and i in split_by_space:
                        str_ = ' ' + buff[i]  + str_
                    else:
                        str_ = buff[i] + str_
            str_list_1[idx] = (str_, POS)
    # print('str_list',str_list_1)
    # delete connected term in related_list
    idx = len(str_list_1)-1
    while idx > 0:        
        if str_list_1[idx][1] in related_list:
            pre_POS_idx = related_list.index(str_list_1[idx][1])
            while str_list_1[idx-1][1] in related_list:
                if related_list.index(str_list_1[idx-1][1]) < pre_POS_idx:
                    pre_POS_idx = related_list.index(str_list_1[idx-1][1])
                    # print('q',str_list_1[idx-1])
                    del str_list_1[idx-1]
                    idx -= 1
                    # if str_list_1[idx-1][1] in related_list:
                    #     print('qc',str_list_1[idx-1])
                    #     pre_POS_idx = related_list.index(str_list_1[idx-1][1])
                else:
                    break
        idx -= 1
    return str_list_1

def merge_with_order(str_list, related_list):
    str_list_1 = str_list[:]
    idx = 0
    while idx <len(str_list_1):
        if str_list_1[idx][1] in related_list:            
            str_buf = str_list_1[idx][0]
            while idx+1 < len(str_list_1):            
                if str_list_1[idx+1][1] in related_list:
                    str_buf += str_list_1[idx+1][0]
                    del str_list_1[idx+1]
                else:
                    break
            str_list_1[idx:idx+1] = [(str_buf, str_list_1[idx][1])]
        idx += 1
    return str_list_1
#mark same POS (place name) by 、
# def same_POS(str_list_1):
#     idx = 0
#     while idx < len(str_list_1):
#         POS_, str_ = str_list_1[idx]

# mark start time and end time
def remark_POS_with_re(str_list, mapping_str, POS_list):
    str_list_1 = str_list[:]
    mapping_split = re.split('(<((?!>).)+>)', mapping_str)
    idx = 0
    while idx < len(mapping_split):
        if re.match('<.+>', mapping_split[idx]) is not None:
            del mapping_split[idx+1]
        idx += 1
    if DEBUG:
        print('map',mapping_split)
    mapping_split = [i for i in mapping_split if len(i)!=0]
    num_split = len(mapping_split)
    check_point = [0 for i in range(len(mapping_split))]
    if DEBUG:
        print('chk',check_point)
    # POS_loc = []
    for idx in range(len(mapping_split)):
        if '<' in mapping_split[idx]:
            check_point[idx] = 1
            # POS_loc.append(idx)
    # print(POS_loc, check_point)
    # print('rr',len(str_list_1)-num_split+1, num_split)
    # print(mapping_split)
    # print(POS_loc,check_point)
    for idx in range(len(str_list_1)-num_split+1):
        str_ = str_list_1[idx: (idx+num_split)]
        # print('ge',str_)
        map_target = ''
        # cmb_str = [str_p[chk] for i in zip(str_, check_point)]
        for str_p, chk in zip(str_, check_point):
            # print(str_p, chk)
            if chk == 0:
                map_target += str_p[chk]
            else:
                map_target += '<'+ str_p[chk] + '>'
        # print('dd',map_target, mapping_str)
        if re.search(mapping_str, map_target) is not None: #pattern match
            for idx_ in range(num_split):
                if check_point[idx_] == 0 and idx_ == 0: #first pattern with POS S
                    if mapping_split[idx_] == str_list_1[idx+idx_][0]:
                        str_list_1[idx+idx_:idx+idx_+1] = [(mapping_split[idx_], POS_list[idx_])]
                    else:
                        pre_str, get_str = str_list_1[idx+idx_][0][:-len(mapping_split[idx_])], mapping_split[idx_]
                        pre_POS, get_POS = str_list_1[idx+idx_][1], POS_list[idx_]
                        str_list_1[idx+idx_:idx+idx_+1] = [(pre_str,pre_POS),(get_str, get_POS)]
                elif check_point[idx_] == 0 and idx_ == num_split-1:
                    if mapping_split[idx_] == str_list_1[idx+idx_][0]:
                        str_list_1[idx+idx_:idx+idx_+1] = [(mapping_split[idx_], POS_list[idx_])]
                    else:
                        get_str, post_str = mapping_split[idx_], str_list_1[idx+idx_][0][len(mapping_split[idx_]):] 
                        get_POS, post_POS = POS_list[idx_], str_list_1[idx+idx_][1]
                        str_list_1[idx+idx_:idx+idx_+1] = [(get_str, get_POS), (post_str,post_POS)]
                else:
                    str_list_1[idx+idx_:idx+idx_+1] = [(str_list_1[idx+idx_][0], POS_list[idx_])]
            # for (idx_, rePOS) in zip(POS_loc, POS_list):
            #     str_list_1[idx+idx_] = (str_list_1[idx+idx_][0], rePOS)
    return str_list_1
            
def remark_POS(str_list, POS_list, new_POS):
    str_list_1 = str_list[:]
    for idx in range(len(str_list_1)):
        if str_list_1[idx][1] in POS_list:
            str_list_1[idx] = (str_list_1[idx][0], new_POS)
    return str_list_1

def news_rel_evalu(str_list, pos_keyword_list, neg_keyword_list):
    pos_v = 1
    neg_v = -90
    evalu = 0
    for pos_keyword in pos_keyword_list:
        for strp in str_list:
            if strp[1] == 'S' and pos_keyword in strp[0]:
                evalu += pos_v
                continue
    for neg_keyword in neg_keyword_list:
        for strp in str_list:
            if strp[1] == 'S' and neg_keyword in strp[0]:
                evalu += neg_v
                continue
    return evalu

def extract_attr(str_list, POS_list, drop_tuple_trigger = None):
    str_list_1 = str_list[:]  
  
    POS_seq_in_doc = []
    for str_p in str_list_1:
        if str_p[1] in POS_list and str_p[1] not in POS_seq_in_doc:
            POS_seq_in_doc.append(str_p[1])
    # print(POS_seq_in_doc)
    # lost_POS_in_doc = [i for i in POS_list if i not in POS_seq_in_doc]
    extract_all = []
    current_state = ['' for i in POS_seq_in_doc]
    first_pair_flg = True
    extract = []
    pre_state = None
    # print('extractSS',str_list_1)
    for idx, str_p in enumerate(str_list_1):
        pre_state = current_state
        # print('1 all',extract_all)
        if str_p[1] in POS_seq_in_doc:
            c_idx = POS_seq_in_doc.index(str_p[1])
            # handle for the new currenet_state 
            if first_pair_flg == False:                
                for idx_c in range(c_idx, len(current_state)):
                    current_state[idx_c] = ''
                current_state[c_idx] = str_p[0]
                # print('aa',str_p[0])
                first_pair_flg = True
            # handle pair extraction     
            if c_idx == len(POS_seq_in_doc)-1: #information complete
                extract.append(current_state[:])
                for _idx in range(len(extract)):
                    extract[_idx][-1] = str_p[0]
                extract_all += extract
                extract = []
                first_pair_flg = False
            else:
                if current_state[c_idx] != '':
                    extract.append(current_state[:])
                    # extract += [current_state[:]]
                current_state[c_idx] = str_p[0]
                for _idx in range(len(extract)):
                    if extract[_idx][c_idx] == '':
                        extract[_idx][c_idx] = str_p[0]
        if str_p[1] == 'S' and '才' in str_p[0]:            
            # print('才s', extract)
            current_state = pre_state
            extract = []
    # re-order attr
    # print(extract_all)
    re_order = []
    for POS in POS_list:
        if POS in POS_seq_in_doc:
            re_order.append(POS_seq_in_doc.index(POS))
        else:
            re_order.append('None')
    re_order_extract = []
    for _it in extract_all:
        ext = []
        for idx_r in re_order:
            if idx_r == 'None':
                ext.append('None')
            else:
                ext.append(_it[idx_r])
        re_order_extract.append(ext)    
    return re_order_extract

def output_docx(document, title, attrs, column_list):
    document.add_paragraph('\n'+title)
    table = document.add_table(rows=1, cols=len(column_list)+1, style='Light Grid Accent 1')
    hdr_cells = table.rows[0].cells
    for idx in range(len(column_list)):
        hdr_cells[idx+1].text = column_list[idx]

    if attrs is None or len(attrs) == 0:
        row_cells = table.add_row().cells
        row_cells[0].text = '機器產生'
        row_cells = table.add_row().cells
        row_cells[0].text = '人工產生'    
    else:
        if len(attrs)!=0:
            # print('坡地災害-道路中斷:')
            for vec in attrs:
                row_cells = table.add_row().cells
                row_cells[0].text = '機器產生'
                for idx in range(len(vec)):
                    row_cells[idx+1].text = vec[idx]
                row_cells = table.add_row().cells
                row_cells[0].text = '人工產生'

def time_Handle(str_list, announce_time):
    str_list_1 = str_list[:] 
    def week_init(time):
        week_day = (Time.localtime(time).tm_wday+1)%7
        base_time = time-week_day*60*60*24
        return base_time
    
    def cht_nums_to_num(cht_nums):     #cht_nums less than 99   
        n_cht = ['一','二','三','四','五','六','七','八','九']
        num = 0
        if '十' in cht_nums:
            nn = cht_nums.split('十')
            if nn[0] == '':
                num += 10
            else:
                num += n_cht.index(nn[0])*10

            if nn[1] != '':
                num += n_cht.index(nn[1])            
        else:
            num = n_cht.index(cht_nums)
        return num
            

    # w_days= ['[周週]日','[周週]一', '[周週]二', '[周週]三', '[周週]四', '[周週]五', '[周週]六']
    w_days= ['日', '一', '二', '三', '四', '五', '六']
    des_day = None    
    #replace time string
    for idx in range(len(str_list_1)):
        if str_list_1[idx][1] == 'Day2':
            # print(str_list_1[idx][0])
            if re.search('[上下]?[周週][一二三四五六日]', str_list_1[idx][0]):
                mult = 0
                if str_list_1[idx][0][0] == '上':
                    mult = -1
                elif str_list_1[idx][0][0] == '下':
                    mult = 1            
                basetime = week_init(announce_time + mult*7*60*60*24)

                w_day = w_days.index([i for i in w_days if i in str_list_1[idx][0]][0])
                des_day = basetime + w_day*60*60*24
                # for cond_idx in range(len(w_days)):
                #     if re.search(w_days[cond_idx], str_list_1[idx][0]) is not None:
                #         des_day = basetime + cond_idx*60*60*24
                #         break
                str_list_1[idx] = (Time.strftime("%Y-%m-%d", Time.localtime(des_day)), str_list_1[idx][1])
                # return Time.strftime("%Y-%m-%d", des_day)
            elif re.search('[今昨][天日]{0,1}', str_list_1[idx][0]) is not None:
                des_day = announce_time
                if '昨' in str_list_1[idx][0]:
                    des_day = announce_time - 60*60*24
                str_list_1[idx] = (Time.strftime("%Y-%m-%d", Time.localtime(des_day)), str_list_1[idx][1])
            elif re.search('前天', str_list_1[idx][0]) is not None:                                
                des_day = announce_time - 60*60*24*2
                str_list_1[idx] = (Time.strftime("%Y-%m-%d", Time.localtime(des_day)), str_list_1[idx][1])
        elif str_list_1[idx][1] == 'Day':
            # print(str_list_1[idx][0])
            if re.match('[0-9一二三四五六七八九十]+日',str_list_1[idx][0]) is not None:
                tmp_day_str = Time.strftime("%Y-%m-%d", Time.localtime(announce_time)).split('-')
                dd = None
                if re.match('[一二三四五六七八九十]+日',str_list_1[idx][0]) is not None:
                    dd = str(cht_nums_to_num(str_list_1[idx][0][:-1]))
                else:
                    dd = str(str_list_1[idx][0][:-1])
                if len(dd) == 1:
                    dd = '0'+dd
                tmp_day_str[2] = dd
                str_list_1[idx] = ('-'.join(tmp_day_str), str_list[idx][1])
            elif re.match('[0-9]{4}年[0-9]{1,2}月[0-9]{1,2}日',str_list_1[idx][0]) is not None:
                patt = re.match('([0-9]{4})年([0-9]{1,2})月([0-9]{1,2})日',str_list_1[idx][0])
                str_list_1[idx] = ('-'.join([patt.group(1), patt.group(2), patt.group(3)]), str_list[idx][1])
            elif re.match('[0-9]{2,3}年[0-9]{1,2}月[0-9]{1,2}日',str_list_1[idx][0]) is not None:
                patt = re.match('([0-9]{2,3})年([0-9]{1,2})月([0-9]{1,2})日',str_list_1[idx][0])
                year = str(int(patt.group(1))+1911)
                str_list_1[idx] = ('-'.join([year, patt.group(2), patt.group(3)]), str_list[idx][1])
            elif re.match('[0-9]{1,2}月[0-9]{1,2}日',str_list_1[idx][0]) is not None:
                year = Time.strftime("%Y-%m-%d", Time.localtime(announce_time)).split('-')[0]
                patt = re.match('([0-9]{1,2})月([0-9]{1,2})日',str_list_1[idx][0])
                str_list_1[idx] = ('-'.join([year, patt.group(1), patt.group(2)]), str_list[idx][1])
        elif str_list_1[idx][1] == 'time3':
            ctime = int(re.search('[早晚]([0-9])+[時點]', str_list_1[idx][0]).group(1))
            if '晚' in str_list_1[idx][0]:
                ctime += 12
            str_list_1[idx] = (str(ctime)+'時', str_list[idx][1])
    # # replace
    # for idx in range(len(str_list_1)):

    return str_list_1    



def parser(strA, announce_time):
    attr_tables = []
    # strA = [(u'坡地災害部分高雄市燕巢區東燕里中民路61巷邊坡坍方，導致一戶民宅內3名民眾遭到掩埋，以及高雄市旗山區新光里後厝巷民宅後方崩塌，導致土石衝破房屋圍牆，造成民宅受損最為嚴重。','S')]
    # print('dd',strA)
    strA = filter_re(strA, '，', 'Prep')
    strA = filter(strA, [place_dict, road_dict])
    strA = filter_re(strA, '[北中南]橫公路', 'Road')
    strA = filter_re(strA, '[0-9一二三四五六七八九]+年', 'Year')
    strA = filter_re(strA, '[0-9一二三四五六七八九十廿]+月', 'Month')
    strA = filter_re(strA, '[0-9一二三四五六七八九十廿]+日', 'Day')
    strA = filter_re(strA, '[台投苗北屏][0-9一二三四五六七八九十廿]+[甲乙丙丁戊己庚辛]{0,1}線', 'Road')
    strA = filter_re(strA, '[台投苗北屏]{0,1}[0-9一二三四五六七八九十廿]+[甲乙丙丁戊己庚辛]{0,1}公路', 'Road')
    strA = filter_re(strA, '北宜公路', 'Road')
    strA = filter_re(strA, '蘇花公路', 'Road')
    strA = filter_re(strA, r'[0-9\.,]+km2', 'unit')
    strA = filter_re(strA, r'[0-9\.,]+km3', 'unit')
    # strA = filter_re(strA, '[0-9]+巷', 'addr0')
    # strA = filter_re(strA, '..巷', 'addr0')
    # strA = filter_pre_re(strA, ['Road'], '[一二三四五六七八九十]+段', 'addr1')
    # strA = filter_pre_re(strA, ['Road','addr0'], '[0-9]+[巷]', 'addr2')
    # strA = filter_pre_re(strA, ['Road','addr0', 'addr1', 'addr2'], '[0-9]+[弄]', 'addr3')
    # strA = filter_pre_re(strA, ['Road','addr0', 'addr1', 'addr2', 'addr3'], '[0-9]+[號]', 'addr4')
    for re_ in ['早上', '上午', '中午', '下午', '晚上', '傍晚', '晚間', '凌晨']:
        strA = filter_re(strA, re_, 'time1')
    for re_ in ['[上下]?[周週][一二三四五六日]', '[今昨][天日]{0,1}', '前天']:
        strA = filter_re(strA, re_, 'Day2')
    #test
    # strA = [('公路總局', 'S'), ('昨', 'Day2'), ('封閉', 'S'), ('台9線北宜公路', 'Road'), ('坪林', 'Town'), ('到', 'S'), ('頭城', 'Town'), ('和', 'S'), ('台7線北橫公路', 'Road'), ('羅浮到明池路段', 'S'), ('，', 'Prep'), ('台9線', 'Road'), ('昨', 'Day2'), ('晚7時因坍方', 'S'), ('，', 'Prep'), ('晚間', 'time1'), ('緊急封路', 'S')]
    # print('time_3_1',strA)
    strA = filter_pre_re(strA, ['Day', 'Day2'], '[早晚][0-9]+[時點]', 'time3') #new

    # print('time_3_2',strA)
    # raw_input( )
    
    # strA = filter_re(strA, re_, 'Day2')
    strA = filter_pre_re(strA, ['Year', 'Month', 'Day', 'time1'], r'[0-9]+[時點][0-9]+分?|[0-9]+[時點]', 'time2') #一二三四五六七八九十
    # print('1',strA)
    # strA = merge_term(strA, ['County','Town','Village', 'Road', 'addr0', 'addr1', 'addr2', 'addr3', 'addr4'])
    strA = merge_term(strA, ['County','Town','Village'])
    strA = merge_term(strA, ['Year', 'Month','Day'])
    # print('tt', announce_time)
    # print('time1', strA)
    strA = time_Handle(strA, announce_time)
    # print('time2', strA)
    strA = merge_term(strA, ['Year', 'Month','Day','time1', 'time2'], [3])
    # print('time3', strA)
    # print('1.5',strA)
    strA = merge_term(strA, ['Day2','time1', 'time2'])
    # print('time4', strA)
    strA = merge_term(strA, ['Day2', 'time3'], [1])
    # print('time5', strA)
    strA = remark_POS(strA, ['Year', 'Month','Day','Day2','time1', 'time2', 'time3'], 'time')
    strA = remark_POS(strA, ['County','Town','Village', 'addr0', 'addr1', 'addr2', 'addr3', 'addr4'], 'Place')
    
    idx = 0
    while idx < len(strA)-1:
        if strA[idx][1] == 'Place' and strA[idx+1][1] == 'S' and strA[idx+1][0][0] == '橋':
            strA[idx:idx+2] = [(strA[idx][0]+strA[idx+1][0], 'S')]
        idx += 1
    #find start time and end time
    # strA = remark_POS_with_re(strA, '<time>止',['end_time'])
    strA = remark_POS_with_re(strA, '<time>[至到]<time>',['start_time', 'S','end_time'])
    strA = remark_POS_with_re(strA, '<time>發生',['start_time', 'S'])
    strA = remark_POS(strA, 'time', 'start_time')
    # strA = remark_POS(strA, ['time'], 'start_time')
    # print(strA)
    # print('2',strA)
    # 坡地災害-道路中斷
    strA_R = filter_re(strA, '[0-9一二三四五六七八九十廿百千萬點\.,]+[至到][0-9一二三四五六七八九十廿百千萬點\.,]+公里', 'unit', replace_re_rule= [('[至到]', 'k-'), ('公里','k')])
    # print('3',strA_R)
    strA_R = filter_re(strA_R, '[0-9一二三四五六七八九十廿百千萬點\.,]+[至到][0-9一二三四五六七八九十廿百千萬點\.,]+[kK]', 'unit', replace_re_rule= [('[至到]', 'k-'), ('公里','k'), ('K','k')])
    # print('4',strA_R)
    strA_R = filter_re(strA_R, '[0-9一二三四五六七八九十廿百千萬點\.,]+[kK][至到][0-9一二三四五六七八九十廿百千萬點\.,]+[kK]', 'unit', replace_re_rule= [('[至到]', '-'), ('K','k')])
    # print('5',strA_R)
    # print(strA_R)
    strA_R = filter_re(strA_R, '[0-9一二三四五六七八九十廿百千萬點\.,]+公里[至到][0-9一二三四五六七八九十廿百千萬點\.,]+公里', 'unit', replace_re_rule= [('[至到]', '-'),('公里','k')])
    # print('6',strA_R)
    strA_R = filter_re(strA_R, '[0-9一二三四五六七八九十廿百千萬點\.,]+公里路段', 'unit3')
    strA_R = filter_re(strA_R, '路基', 'temp')
    strA_R = filter_re(strA_R, '[0-9一二三四五六七八九十廿百千萬點\.,]+公里', 'unit', 'unit3')    
    strA_R = filter_re(strA_R, '[0-9一二三四五六七八九十廿百千萬點\.,]+公里路', 'unit2')
    strA_R = filter_re(strA_R, '[0-9一二三四五六七八九十廿百千萬點\.,]+公里', 'unit')  
    strA_R = filter_re(strA_R, '[0-9一二三四五六七八九十廿百千萬點\.,]+km', 'unit')
    strA_R = filter_re(strA_R, '[0-9一二三四五六七八九十廿百千萬點\.,]+[kK]', 'unit')
    strA_R = remark_POS_with_re(strA_R, '<Place>[至到]<Place>路段', ['unit_', 'unit_', 'unit_', 'unit_'])    
    # print('R_end1',strA_R)
    strA_R = merge_with_order(strA_R, ['unit_'])
    strA_R = remark_POS(strA_R, 'unit_', 'unit')
    if DEBUG:
        print('R_end2',strA_R)
    idx = 0
    road_merge_flg = False
    while idx < len(strA_R):
        tup = strA_R[idx]
        if tup[1] == 'Road':
            substrA = strA_R[idx+2:idx+5]

            # print('LL', substrA)

            specical_flg= False
            for ss in substrA:                
                if re.match('路?段', ss[0]) is not None and ss[1] == 'S':
                    specical_flg = True                    
                    break
            if specical_flg:
                # print('LG', strA_R, 'LL', substrA)
                road_merge_flg = True
        idx += 1
        if road_merge_flg:
            # strA_R[idx]
            # print('ttl',strA_R[idx])
            while strA_R[idx+1][1] in ['S', 'County', 'Town', 'Village', 'Place']:
                strA_R[idx] = (strA_R[idx][0] + strA_R[idx+1][0], 'S')
                # print('ttl',strA_R[idx])
                if '段' in strA_R[idx+1][0] and strA_R[idx+1][1] == 'S':
                    strA_R[idx+1:idx+2] = []                    
                    break
                strA_R[idx+1:idx+2] = []
            # print('LG', strA_R)
            road_merge_flg = False


    strA_R = filter_pre_re(strA_R, ['Road'], '.{2,10}路段', 'unit')
    strA_R = filter_pre_re(strA_R, ['Road'], '.{2,10}段', 'unit')

    idx = 0
    while idx < len(strA_R)-1:
        if strA_R[idx][1] == strA_R[idx+1][1]:
            strA_R[idx] = (strA_R[idx][0]+strA_R[idx+1][0], strA_R[idx][1])
            del strA_R[idx+1]
        idx += 1
    strA_R = [i for i in strA_R if i[1] in ['Road', 'unit', 'Place', 'start_time', 'end_time']]
    if DEBUG:
        print('dd1',strA_R)
    return strA_R
    # # print(strA_R)
    # # if news_rel_evalu(strA_R, ['道路中斷', '交通中斷', '封路', '交通阻斷', '阻斷交通', '封閉路線', '土石坍方中斷', '封閉', '路基流失中斷','坍方中斷'], []):
    # pos_list = ['道路中斷' ,'交通中斷', '封路', '交通阻斷', '阻斷交通', '封閉路線', '土石坍方中斷', '路基流失中斷','坍方中斷', '路基沖刷', '路基流失', '路基嚴重受損', '對外交通中斷', '道路坍方']
    # neg_list = ['預警性封閉', '暫時性封閉', '施工', '高乘載管制', '工程', '高承載管制', '車潮', '921大地震', '921地震']
    # eval_ = 1 # news_rel_evalu(strA_R, pos_list, neg_list)    
    # if eval_> 0:
    #     strA1 = extract_attr(strA_R, ['Road', 'unit', 'Place', 'start_time', 'end_time'])
    # else: 
    #     strA1 = None
    # # print('坡地災害-道路中斷:', eval_, strA1)
    # # strA1 = extract_attr(strA, ['Road', 'Place', 'start_time', 'end_time'], ['道路中斷', '交通中斷', '封路', '交通阻斷', '阻斷交通', '封閉路線', '土石坍方中斷', '封閉'])
    # #docx print
    # attr_tables.append(strA1)
    # # output_docx(document, '坡地災害-道路中斷:', strA1, ['道路', '里程數', '地區', '起始時間', '截止時間'])    

    # # strA1 = strA
    # # for unit in ['點','筆','處']:        
    # #     strA1 = filter_re(strA1, r'[0-9一二三四五六七八九十百千萬點\.,]+'+unit, 'unit')
    # #     # strA1 = filter_re(strA1, r'[一二三四五六七八九十百千萬]+'+unit, 'unit')
    # # # print(strA1)
    # # eval_ = 1 # news_rel_evalu(strA1, ['坡地災點','坡地災害','坡地災情'], neg_list)
    # # if eval_>0:
    # #     strA1 = extract_attr(strA1, ['Place', 'unit', 'start_time', 'end_time'])
    # # else: 
    # #     strA1 = None    
    # # # if strA1 is not None:
    # # #     if len(strA1)!=0:
    # # #         print('坡地災害-災點數:',strA1)
    # # attr_tables.append(strA1)
    # # # output_docx(document, '坡地災害-災點數:', strA1, ['地點範圍', '坡地災點數', '起始時間', '截止時間'])    
                
    # # strA1 = strA
    # # for unit in ['立方公尺','m3']:        
    # #     strA1 = filter_re(strA1, r'[0-9一二三四五六七八九十百千萬點\.,]+'+unit, 'unit')
    # #     # strA1 = filter_re(strA1, r'[一二三四五六七八九十百千萬]+'+unit, 'unit')
    # # eval_ = 1 # news_rel_evalu(strA1, ['土方量','土石量','土方總量','坍方量'], neg_list)
    # # if eval_ > 0:
    # #     strA1 = extract_attr(strA1, ['Place', 'unit', 'time'])
    # # else: 
    # #     strA1 = None    
    # # attr_tables.append(strA1)
    # # # output_docx(document, '坡地災害-土方量:', strA1, ['地點範圍', '土方量 (立方公尺)', '時間'])
    # return attr_tables

def comp_entry(e1, e2, mech_arr):
    def checkRelName(str1, str2):
        for char in str1:
            if char in str2:
                str2 = str2.replace(char,'',1)
                str1 = str1.replace(char,'',1)
        if len(str1) == 0 or re.match('[鄉鎮縣市村里區]+', str2).span()[1] == len(str1) or len(str2) == 0 or re.match('[鄉鎮縣市村里區]+', str2).span()[1] == len(str2):
            return True
        else:
            return False
    
    def checkRelTime(str1, str2):
        t1, t2 = Time.mktime(Time.strptime(str1,"%Y-%m-%d")), Time.mktime(Time.strptime(str2,"%Y-%m-%d"))
        if abs(t1-t2)<=86400:
            return True
        else:
            return False

    rel_list = []        
    for item1, item2, mech in zip(e1,e2,mech_arr):
        if mech == 'Name':
            rel_list.append(checkRelName(item1, item2))
        elif mech == 'Time':
            rel_list.append(checkRelTime(item1, item2))



def parser2(strA, announce_time):
    neg_keywords = ['定時性?', '預警性?','預防性?' , '間歇性?', '高[乘承]載', '將', '視', '易發生', '考量', '結冰', '積雪', '蟹', '演練', '演習', '拒馬', '拓寬', '拆除', '警戒區', '不排除', '恢復','未達封路標準','去年', '前年', '往年', '每當', '每日', '每逢']
    announce_time = Time.mktime(Time.strptime(announce_time, "%Y-%m-%d"))
    # strA = [(u'截至8月1日止，蒐整公路總局及農委會水土保持局之坡地災點共41點，主要為道路中斷災情，分佈在台7線、台8線、台9線、台18線、台20線、台21線、台26線等處，災點分布如圖15所示。','S')]
    # document.add_paragraph(strA[0][0])
    # print(strA[0][0])
    # print('extract:')
    def contian_negKeyword(neg_keywords, text):
        flg = False
        for it_ in neg_keywords:
            if re.search(it_, text) is not None:
                flg = True
                break
        return flg

    strs = re.split('[,，;。!！；]',strA[0][0])
    if DEBUG:
        print('q1',strs,'\n')
    strs = [[(i,'S')] for i in strs if len(i)!=0 and not contian_negKeyword(neg_keywords, i) and len(i)!=0]
    if strs == []:
        strs = [[(' ','S')]]
    if DEBUG:
        print('q2',strs,'\n')
    # detail news will be extrated
    # strs.append(strA)   #new

    extract_items = []
    for idx, string in enumerate(strs):
        # print('subnews',string)
        extract_items += parser(string, announce_time)
        # attr_tables = parser(string, announce_time)
        # if idx == 0:
        #     tables = attr_tables
        # else:
        #     for t_idx in range(len(tables)):
        #         tables[t_idx] = tables[t_idx] + attr_tables[t_idx]
    if DEBUG:
        print('tab1', extract_items)
    table = extract_attr(extract_items, ['Road', 'unit', 'Place', 'start_time', 'end_time'])
    if DEBUG:
        print(table)
    # def n_element_enty(entry):
    #     n_ele = 0
    #     for term in entry:
    #         if term not in ['','None']:
    #             n_ele +=1
    #     return n_ele


    # #過濾table 中element數量較少的entry
    # n_element = []
    # if FILTER_FEW_ELEMENT:
    #     for table in tables:
    #         max_n = 0
    #         for entry in table:
    #             n_ele = n_element_enty(entry)
    #             if max_n < n_ele:
    #                 max_n = n_ele
    #         n_element.append(max_n)
    #     for t_idx in range(len(tables)):
    #         r_idx = len(tables[t_idx])-1
    #         while r_idx >= 0:
    #             if n_element_enty(tables[t_idx][r_idx]) != n_element[t_idx]:
    #                 del tables[t_idx][r_idx]
    #             r_idx += -1
    # remove residual rows
    merge_idx = 0
    important_idx = 0 # 個別table必定要有的element，缺的row移除 
    # print('tb_b', tables)
    # for t_idx in range(len(tables)):
    r_idx = len(table)-1
    while r_idx >= 0:
        # chk_idx = [i for i in range(r_idx+1, len(table)) if i != merge_idx] #modifyyyyyyyyyyyyyyyyyyyyyyyy            
        chk = len([1 for i in range(0, r_idx) if smaller(table[r_idx], table[i], merge_idx)])
        # chk2 = len([1 for i in range(r_idx+1, len(table)) if table])
        if chk>0 or table[r_idx][important_idx] in ['', 'None']:
            del table[r_idx]
        r_idx += -1
    # print('tb_r', table)
    return table
    # print()

if __name__ == '__main__':
    # with open('news.txt') as f:
    #     new_text = re.sub('=','',f.read())
        # new_text = re.sub('\n','',new_text)
        #全形轉半形
        # with open("坡地災害-道路中斷.csv", "w") as f:
        #     f.write('\n')
        # new_text = new_text.translate(str.maketrans('０１２３４５６７８９　', '0123456789 '))
        # new_text = re.sub(' ','',new_text)
        # texts = new_text.split('。')
        # texts = [text+'。' for text in texts if len(text)>0]
        # for t_id, text in enumerate(texts):
        #     if t_id>-1:
        #     # print(text)
        #     # print()
        #         parser2([(text,'S')])

    neg_keywords = ['定時性?', '預警性?','預防性?' , '間歇性?', '高[乘承]載', '將', '視', '易發生', '考量', '結冰', '積雪', '蟹', '演練', '演習', '拒馬', '拓寬', '拆除', '警戒區', '不排除', '恢復','未達封路標準','去年', '前年', '往年', '每當', '每日', '每逢']

    document = Document()
    
    directory = './data/'
    files = os.listdir(directory)
    files.sort()
    files = [i for i in files if 'roadclosed' in i]
    # print(files)
    # global announce_time
    news_idx = -1
    for idx, file in enumerate(files):
        news = json.load(open(directory+file, 'r'))
        for _dict in news:
            news_idx += 1
            if news_idx < -5:
                continue
            elif news_idx > 12000:
                break
            if news_idx%100==0:
                print('doc: ',news_idx,'\n')

            text = _dict['text']
            text = text.translate(str.maketrans('０１２３４５６７８９　', '0123456789 '))
            text = re.sub(' ','',text)
            document.add_paragraph('########################################################################################################################################################\n')
            document.add_paragraph('新聞來源:'+_dict['media']+'\n')
            document.add_paragraph('發布時間:'+_dict['time']+'\n')
            document.add_paragraph('URL:'+_dict['url']+'\n')
            document.add_paragraph(text+'\n')
            parser2(neg_keywords,[(text,'S')], announce_time)
            break
            # try:
            #     parser2([(text,'S')])
            # except:
            #     print('error file:', file)
            #     print('error', _dict)
        break
    document.save('demo.docx')
    
        


